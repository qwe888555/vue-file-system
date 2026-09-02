# -*- coding: utf-8 -*-
"""
NeuHub 资源系统 - 管理员端说明手册生成器
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# 截图目录
SHOT_DIR = r'C:\Users\ASUS\.trae-cn\trae-browser-screenshots\6a968bc2127aabb8a5fb5f38'

# 截图文件映射
SHOTS = {
    'knowledge': 'shot-20260901-104012-625606700.jpg',
    'chat': 'shot-20260901-104031-333074500.jpg',
    'faq': 'shot-20260901-104045-740061000.jpg',
    'account': 'shot-20260901-104055-567406400.jpg',
    'department': 'shot-20260901-104105-620077400.jpg',
    'log_overview': 'shot-20260901-104115-103489500.jpg',
    'log_upload': 'shot-20260901-104124-097013400.jpg',
    'log_query': 'shot-20260901-104131-914845800.jpg',
    'log_sensitive': 'shot-20260901-104139-906273800.jpg',
    'log_login': 'shot-20260901-104148-961190500.jpg',
    'log_operation': 'shot-20260901-104157-441631400.jpg',
}

# ── 样式设置 ──
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 设置页边距
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)


def add_title(text, level=0):
    """添加标题"""
    if level == 0:
        p = doc.add_heading(text, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_heading(text, level=level)


def add_para(text, bold=False, indent=False):
    """添加段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p


def add_bullet(text):
    """添加项目符号"""
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p


def add_table(headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(10)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()  # 表后空行


def add_screenshot(key, caption=''):
    """插入截图"""
    if key in SHOTS:
        img_path = os.path.join(SHOT_DIR, SHOTS[key])
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if caption:
                cap_p = doc.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap_p.add_run(f'图：{caption}')
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 128, 128)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            doc.add_paragraph()


def add_video_placeholder():
    """添加演示视频占位符"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('【此处插入演示视频】')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)


def add_page_break():
    """添加分页符"""
    doc.add_page_break()


# ════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════
add_page_break()
for _ in range(6):
    doc.add_paragraph()
add_title('NeuHub 资源系统', level=0)
add_title('管理员端操作说明手册', level=0)
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('成都东软学院')
run.font.size = Pt(14)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('版本：V1.0')
run2.font.size = Pt(12)
run2.font.name = '微软雅黑'
run2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('日期：2026年9月')
run3.font.size = Pt(12)
run3.font.name = '微软雅黑'
run3._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ════════════════════════════════════════════
# 目录说明
# ════════════════════════════════════════════
add_page_break()
add_title('目录', level=1)
toc_items = [
    '1. 系统概述',
    '2. 登录与界面总览',
    '3. 知识库管理',
    '4. 教研问答',
    '5. FAQ 管理',
    '6. 账号管理',
    '7. 部门管理',
    '8. 日志管理',
    '9. 角色权限说明',
    '附录',
]
for item in toc_items:
    add_para(item)

# ════════════════════════════════════════════
# 1. 系统概述
# ════════════════════════════════════════════
add_page_break()
add_title('1. 系统概述', level=1)
add_para(
    'NeuHub 资源系统是成都东软学院一站式智能知识库系统，基于 RAG（检索增强生成）技术，'
    '为全校师生提供智能问答、知识库检索、文档管理等服务。系统分为管理员端和用户端，'
    '本手册面向管理员用户，详细说明管理员端各功能模块的操作方法。', indent=True)
add_para('系统主要功能模块包括：', bold=True)
add_bullet('知识库管理：上传、创建、编辑、删除知识文档，支持在线预览')
add_bullet('教研问答：与 AI 进行智能对话，管理对话历史')
add_bullet('FAQ 管理：管理常见问题，支持草稿、发布、驳回状态')
add_bullet('账号管理：管理系统用户账号、角色与所属单位')
add_bullet('部门管理：维护学校组织架构，支持树形层级管理')
add_bullet('日志管理：查看系统操作日志、上传日志、查询日志、敏感内容、登录日志')

# ════════════════════════════════════════════
# 2. 登录与界面总览
# ════════════════════════════════════════════
add_page_break()
add_title('2. 登录与界面总览', level=1)

add_title('2.1 系统登录', level=2)
add_para('访问系统网址 http://8.156.65.242/ ，进入登录页面。', indent=True)
add_para('登录方式：', bold=True)
add_bullet('账号密码登录：输入工号/学号和密码，点击"登录"按钮')
add_bullet('扫码登录：使用统一身份认证扫码登录')
add_bullet('统一身份认证登录：点击跳转至学校统一认证平台')

add_title('2.2 主界面布局', level=2)
add_para(
    '管理员登录后，左侧为导航菜单栏，包含六个功能模块入口；右侧为内容展示区。'
    '页面右上角显示当前登录用户信息（账号、角色）及退出按钮。', indent=True)

add_table(
    ['菜单项', '功能说明', '访问路径'],
    [
        ['知识库管理', '管理知识库文档', '/knowledge/list'],
        ['教研问答', 'AI 智能问答对话', '/chat'],
        ['FAQ 管理', '常见问题管理', '/faq-manage'],
        ['账号管理', '用户账号管理', '/admin/users'],
        ['部门管理', '组织架构管理', '/departments'],
        ['日志管理', '系统日志查看', '/logs'],
    ]
)

# ════════════════════════════════════════════
# 3. 知识库管理
# ════════════════════════════════════════════
add_page_break()
add_title('3. 知识库管理', level=1)
add_screenshot('knowledge', '知识库管理页面')
add_para(
    '知识库管理是系统的核心模块，用于管理 AI 问答所依赖的知识库文档。'
    '管理员可以上传各类文档（PDF、Word、TXT、图片、音视频等），或在线创建文档，'
    '上传后系统会自动进行向量化处理，供 AI 检索使用。', indent=True)

add_title('3.1 上传文件', level=2)
add_para('点击"上传文件"标签页，可通过以下方式上传文档：', indent=True)
add_bullet('拖拽上传：将文件直接拖到上传区域')
add_bullet('点击上传：点击上传区域选择文件')
add_para('支持格式：PDF、Word（.docx）、TXT、Markdown（.md）、PPT（.pptx）、图片、音视频等。', bold=True)

add_title('3.2 创建文件', level=2)
add_para('点击"创建文件"标签页，可在线创建 Markdown 格式文档，无需本地编写后再上传。', indent=True)

add_title('3.3 资料列表', level=2)
add_para('资料列表展示所有已上传的文档，包含以下信息：', indent=True)
add_table(
    ['列名', '说明'],
    [
        ['资料名', '文档名称，点击可在线预览内容'],
        ['上传单位', '文档所属学院/部门'],
        ['上传者', '上传该文档的用户'],
        ['上传时间', '文档上传的时间'],
        ['操作', '下载、编辑、删除文档'],
    ]
)

add_title('3.4 文档操作', level=2)
add_para('下载：点击"下载"按钮，将文档下载到本地。', indent=True)
add_para('编辑：点击"编辑"按钮，可修改文档信息（如归属单位、描述等）。', indent=True)
add_para('删除：点击"删除"按钮，可删除该文档（需确认）。', indent=True)
add_para('预览：点击文档名称，弹出预览对话框，在线查看文档内容。'
         'Markdown 文档会渲染为格式化显示（标题、表格、代码高亮等），'
         '与原文档格式一致。', indent=True)

add_title('3.5 搜索与筛选', level=2)
add_para('顶部搜索框支持按资料名、上传单位、上传者、上传时间、资料描述进行模糊搜索。', indent=True)

add_title('3.6 分页', level=2)
add_para('列表底部显示总条数和分页控件，支持切换每页显示条数（8/15/20条）及页码跳转。', indent=True)

add_video_placeholder()

# ════════════════════════════════════════════
# 4. 教研问答
# ════════════════════════════════════════════
add_page_break()
add_title('4. 教研问答', level=1)
add_screenshot('chat', '教研问答页面')
add_para(
    '教研问答是系统的智能对话模块，基于 RAG 技术，AI 会根据知识库内容回答用户问题。'
    '管理员可在此页面与 AI 对话，测试知识库效果，也可管理对话历史。', indent=True)

add_title('4.1 对话界面', level=2)
add_para('左侧为对话历史列表，右侧为对话区域。', indent=True)
add_para('顶部显示"教研问答"标题及"退出问答"按钮。', indent=True)

add_title('4.2 发起对话', level=2)
add_bullet('点击"新建对话"按钮，创建新的对话会话')
add_bullet('在底部输入框输入问题，按回车或点击发送按钮提交')
add_bullet('支持语音输入（点击麦克风图标）')
add_bullet('页面提供快捷问题按钮，点击即可快速提问')

add_title('4.3 对话历史', level=2)
add_bullet('左侧列表显示所有历史对话，按时间倒序排列')
add_bullet('支持按关键词搜索对话历史')
add_bullet('点击历史对话可查看完整对话记录')

add_title('4.4 AI 回复', level=2)
add_bullet('AI 回复支持 Markdown 格式渲染（标题、列表、表格、代码高亮等）')
add_bullet('回复中可能包含引用的知识库文档链接，点击可在线预览文档内容')
add_bullet('文档链接旁设有下载按钮，可直接下载该文档')
add_bullet('每条回复下方有点赞/点踩反馈按钮，用于评估回答质量')

add_video_placeholder()

# ════════════════════════════════════════════
# 5. FAQ 管理
# ════════════════════════════════════════════
add_page_break()
add_title('5. FAQ 管理', level=1)
add_screenshot('faq', 'FAQ 管理页面')
add_para(
    'FAQ 管理用于维护系统的常见问题（Frequently Asked Questions）。'
    '已发布的 FAQ 会在用户端展示，帮助用户快速找到常见问题的答案。', indent=True)

add_title('5.1 状态管理', level=2)
add_para('FAQ 支持三种状态：', indent=True)
add_table(
    ['状态', '说明'],
    [
        ['草稿', '正在编辑，尚未发布，用户端不可见'],
        ['已发布', '已发布到用户端，用户可见'],
        ['已驳回', '审核未通过，需要修改后重新提交'],
    ]
)

add_title('5.2 搜索与筛选', level=2)
add_bullet('按问题关键词搜索')
add_bullet('按分类筛选（全部分类/指定分类）')
add_bullet('按状态筛选（全部/草稿/已发布/已驳回）')
add_bullet('点击"重置"清空所有筛选条件')

add_title('5.3 操作', level=2)
add_bullet('新增 FAQ：创建新的常见问题')
add_bullet('编辑：修改已有 FAQ 的内容、分类、状态等')
add_bullet('发布/驳回：审核 FAQ 并更新状态')
add_bullet('删除：删除不再需要的 FAQ')

add_video_placeholder()

# ════════════════════════════════════════════
# 6. 账号管理
# ════════════════════════════════════════════
add_page_break()
add_title('6. 账号管理', level=1)
add_screenshot('account', '账号管理页面')
add_para(
    '账号管理用于管理系统所有用户账号，包括账号的创建、编辑、删除，'
    '以及角色分配和所属单位关联。', indent=True)

add_title('6.1 账号列表', level=2)
add_para('账号列表展示所有系统用户，包含以下信息：', indent=True)
add_table(
    ['列名', '说明'],
    [
        ['序号', '序号'],
        ['账号', '用户登录账号（工号/学号）'],
        ['角色', '用户角色：超级管理员/管理员/普通用户'],
        ['所属单位', '用户所属学院或部门'],
        ['操作', '编辑、删除账号'],
    ]
)

add_title('6.2 角色说明', level=2)
add_table(
    ['角色', '权限说明'],
    [
        ['超级管理员 (super_admin)', '拥有所有模块的最高权限，可管理一切'],
        ['管理员 (admin/college_admin/dept_admin)', '可管理知识库、FAQ、账号（部分）、查看日志'],
        ['普通用户 (student/teacher)', '仅可使用教研问答和查看已发布 FAQ'],
    ]
)

add_title('6.3 操作', level=2)
add_para('新增账号：点击"+ 新增"按钮，填写账号信息、分配角色、关联所属单位。', indent=True)
add_para('编辑账号：点击"编辑"按钮，修改账号角色、所属单位等信息。', indent=True)
add_para('删除账号：点击"删除"按钮，删除该账号（需确认，谨慎操作）。', indent=True)

add_title('6.4 搜索与筛选', level=2)
add_bullet('按账号关键词搜索')
add_bullet('按角色筛选（全部/普通用户/管理员）')
add_bullet('按所属单位筛选（下拉选择学院/部门）')

add_video_placeholder()

# ════════════════════════════════════════════
# 7. 部门管理
# ════════════════════════════════════════════
add_page_break()
add_title('7. 部门管理', level=1)
add_screenshot('department', '部门管理页面')
add_para(
    '部门管理用于维护学校的组织架构，支持树形层级结构。'
    '部门信息与账号管理、知识库归属等功能关联，是系统基础数据的重要组成部分。', indent=True)

add_title('7.1 部门列表', level=2)
add_para(
    '页面顶部显示统计信息：一级部门数量、二级部门数量。'
    '当前系统共 7 个一级部门、13 个二级部门。', indent=True)
add_para('部门列表采用树形表格展示，包含以下列：', indent=True)
add_table(
    ['列名', '说明'],
    [
        ['部门名称', '部门/学院名称，一级部门可展开/收起查看下级'],
        ['层级', '一级部门 / 二级部门'],
        ['下级部门', '该部门下的子部门数量'],
        ['排序', '显示排序值'],
        ['操作', '删除该部门'],
    ]
)

add_title('7.2 组织架构示例', level=2)
add_bullet('校级领导（一级）→ 党委书记、校长、副校长、党委副书记（二级）')
add_bullet('党政管理部门（一级）→ 党委办公室、党委组织部、党委宣传部等（二级）')
add_bullet('教学管理部门（一级）→ 教务部、教学质量管理与保障部等（二级）')
add_bullet('学生工作部门、科研与服务部门、后勤保障部门、继续教育与国际教育（一级）')

add_title('7.3 操作', level=2)
add_para('刷新：点击"刷新"按钮重新加载部门列表。', indent=True)
add_para('新建部门：点击"+ 新建部门"按钮，填写部门名称、选择层级、上级部门、排序等信息。', indent=True)
add_para('删除：点击对应行的"删除"按钮，删除该部门（需确认）。', indent=True)
add_para('展开/收起：点击一级部门行首的箭头图标，展开或收起下级部门。', indent=True)

add_video_placeholder()

# ════════════════════════════════════════════
# 8. 日志管理
# ════════════════════════════════════════════
add_page_break()
add_title('8. 日志管理', level=1)
add_para(
    '日志管理模块记录系统的各类操作日志，便于管理员追踪系统使用情况、'
    '排查问题、分析数据。包含六个子页面。', indent=True)

add_title('8.1 概览', level=2)
add_screenshot('log_overview', '日志管理 - 概览')
add_para(
    '概览页面以统计卡片形式展示各类日志的汇总数据，'
    '支持按时间范围筛选（今天/本周/本月）。', indent=True)
add_table(
    ['统计项', '说明', '当前数据'],
    [
        ['上传', '文档上传总数及总大小', '总数 10，总大小 2.1MB'],
        ['查询', 'AI 问答查询总数、平均响应时间、点赞率', '总数 56，平均响应 6927ms，点赞率 14.3%'],
        ['敏感内容', '命中敏感词的查询数量', '总数 0'],
        ['登录', '用户登录记录统计', '数据获取失败（接口待修复）'],
        ['操作', '管理员操作记录总数', '总数 6'],
    ]
)

add_title('8.2 上传日志', level=2)
add_screenshot('log_upload', '日志管理 - 上传日志')
add_para('记录所有知识库文档的上传操作，包含以下字段：', indent=True)
add_table(
    ['字段', '说明'],
    [
        ['上传用户', '执行上传操作的用户账号'],
        ['文件名', '上传的文件名称'],
        ['类型', '文件扩展名（docx/md/pdf/pptx 等）'],
        ['大小', '文件大小'],
        ['学院', '文件归属学院'],
        ['状态', '上传结果（success/failed）'],
        ['IP', '上传者 IP 地址'],
        ['时间', '上传时间'],
    ]
)
add_para('支持按文件名、日期范围、状态筛选，共 10 条记录。', indent=True)

add_title('8.3 查询日志', level=2)
add_screenshot('log_query', '日志管理 - 查询日志')
add_para('记录所有用户的 AI 问答查询，包含以下字段：', indent=True)
add_table(
    ['字段', '说明'],
    [
        ['查询用户', '提问的用户账号'],
        ['问题', '用户提出的问题内容'],
        ['命中', '知识库命中数量'],
        ['耗时', 'AI 响应耗时（毫秒）'],
        ['反馈', '用户反馈（赞/踩/无反馈）'],
        ['时间', '查询时间'],
    ]
)
add_para('支持按问题关键词、日期范围筛选，共 20 条记录。', indent=True)

add_title('8.4 敏感内容', level=2)
add_screenshot('log_sensitive', '日志管理 - 敏感内容')
add_para('记录命中敏感词规则的用户查询，便于内容安全管控。', indent=True)
add_table(
    ['字段', '说明'],
    [
        ['用户', '触发敏感内容的用户'],
        ['问题', '用户提出的问题内容'],
        ['匹配规则', '命中的敏感词规则'],
        ['时间', '触发时间'],
    ]
)
add_para('支持按问题关键词、日期范围筛选，当前暂无数据。', indent=True)

add_title('8.5 登录日志', level=2)
add_screenshot('log_login', '日志管理 - 登录日志')
add_para('记录所有用户的登录尝试，包含成功和失败记录：', indent=True)
add_table(
    ['字段', '说明'],
    [
        ['用户名', '登录账号'],
        ['登录方式', '账号密码登录 / SSO Mock 登录'],
        ['结果', '成功 / 失败'],
        ['失败原因', '登录失败时的错误信息'],
        ['IP', '登录者 IP 地址'],
        ['时间', '登录时间'],
    ]
)
add_para('支持按登录方式、结果、日期范围筛选，共 20 条记录。', indent=True)

add_title('8.6 操作日志', level=2)
add_screenshot('log_operation', '日志管理 - 操作日志')
add_para('记录管理员在系统中的关键操作（如新建/删除部门、编辑账号等）：', indent=True)
add_table(
    ['字段', '说明'],
    [
        ['操作人', '执行操作的管理员账号'],
        ['操作', '操作类型（新建/删除/编辑等）'],
        ['对象类型', '操作对象类型（部门/账号/文档等）'],
        ['对象描述', '操作对象的具体描述'],
        ['IP', '操作者 IP 地址'],
        ['时间', '操作时间'],
    ]
)
add_para('支持按操作关键词、日期范围筛选，共 6 条记录。', indent=True)

add_video_placeholder()

# ════════════════════════════════════════════
# 9. 角色权限说明
# ════════════════════════════════════════════
add_page_break()
add_title('9. 角色权限说明', level=1)
add_para(
    '系统采用基于角色的访问控制（RBAC），不同角色拥有不同的功能权限。'
    '管理员端主要面向超级管理员和管理员角色开放。', indent=True)

add_table(
    ['功能模块', '超级管理员', '管理员', '普通用户'],
    [
        ['知识库管理', '全部权限', '全部权限', '—'],
        ['教研问答', '可使用', '可使用', '可使用'],
        ['FAQ 管理', '全部权限', '全部权限', '查看已发布'],
        ['账号管理', '全部权限', '受限（本单位）', '—'],
        ['部门管理', '全部权限', '受限', '—'],
        ['日志管理', '全部权限', '查看', '—'],
    ]
)

add_title('9.1 角色类型', level=2)
add_table(
    ['角色标识', '角色名称', '说明'],
    [
        ['super_admin', '超级管理员', '系统最高权限，可管理所有模块'],
        ['admin', '管理员', '校级管理员，管理范围较广'],
        ['college_admin', '学院管理员', '仅限本学院数据管理'],
        ['dept_admin', '部门管理员', '仅限本部门数据管理'],
        ['student', '学生', '普通用户，仅可问答和查看 FAQ'],
        ['teacher', '教师', '普通用户，仅可问答和查看 FAQ'],
    ]
)

# ════════════════════════════════════════════
# 结尾
# ════════════════════════════════════════════
add_page_break()
add_title('附录', level=1)
add_title('A. 常见问题', level=2)
add_para('Q：上传文档后多久可以被 AI 检索到？', bold=False)
add_para('A：文档上传后系统会自动进行解析和向量化处理，通常在几秒到几十秒内完成，处理完成后即可被 AI 检索。', indent=True)
add_para('Q：Markdown 文档预览格式与原文档不一致怎么办？', bold=False)
add_para('A：系统会自动渲染 Markdown 格式（标题、表格、代码高亮等）。如发现格式异常，请检查源文件是否为标准 Markdown 语法。', indent=True)
add_para('Q：删除文档后 AI 还能回答相关问题吗？', bold=False)
add_para('A：删除文档后，其向量化数据也会被清除，AI 将不再检索到该文档内容。', indent=True)

add_title('B. 联系方式', level=2)
add_para('如遇系统问题，请联系系统管理员。', indent=True)

# ── 保存文档 ──
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'NeuHub资源系统-管理员端操作说明手册.docx')
doc.save(output_path)
print(f'文档已生成：{output_path}')
